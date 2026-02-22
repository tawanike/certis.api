import io
from uuid import UUID

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import select, desc
from sqlalchemy.ext.asyncio import AsyncSession
from fastapi import HTTPException

from src.artifacts.briefs.models import BriefVersion
from src.artifacts.claims.models import ClaimGraphVersion
from src.artifacts.specs.models import SpecVersion
from src.audit.models import AuditEvent, AuditEventType
from src.matter.models import Matter, MatterState
from src.matter.services import MatterService


class ExportService:
    def __init__(self, db: AsyncSession):
        self.db = db

    async def lock_for_export(self, matter_id: UUID, actor_id: UUID) -> dict:
        result = await self.db.execute(
            select(Matter).where(Matter.id == matter_id)
        )
        matter = result.scalar_one_or_none()
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")

        if matter.status != MatterState.QA_COMPLETE:
            raise HTTPException(
                status_code=400,
                detail=f"Matter must be in QA_COMPLETE state to lock. Current: {matter.status.value}",
            )

        # Transition to LOCKED_FOR_EXPORT
        matter_service = MatterService(self.db)
        updated = await matter_service.update_status(matter_id, MatterState.LOCKED_FOR_EXPORT)

        # Log audit event
        event = AuditEvent(
            matter_id=matter_id,
            event_type=AuditEventType.MATTER_LOCKED,
            actor_id=actor_id,
            detail={"previous_status": MatterState.QA_COMPLETE.value},
        )
        self.db.add(event)
        await self.db.commit()

        return updated

    async def generate_docx(self, matter_id: UUID, actor_id: UUID = None) -> bytes:
        result = await self.db.execute(
            select(Matter).where(Matter.id == matter_id)
        )
        matter = result.scalar_one_or_none()
        if not matter:
            raise HTTPException(status_code=404, detail="Matter not found")

        if matter.status != MatterState.LOCKED_FOR_EXPORT:
            raise HTTPException(
                status_code=400,
                detail="Matter must be locked for export before generating DOCX",
            )

        # Fetch authoritative artifacts
        brief = await self._get_authoritative(BriefVersion, matter_id)
        claims = await self._get_authoritative(ClaimGraphVersion, matter_id)
        spec = await self._get_authoritative(SpecVersion, matter_id)

        # Build DOCX
        doc = DocxDocument()

        # -- Title Page --
        self._add_title_page(doc, matter, brief)

        # -- Claims Section --
        self._add_claims_section(doc, claims)

        # -- Specification Sections --
        self._add_spec_sections(doc, spec)

        # -- Abstract --
        self._add_abstract(doc, spec)

        # Serialize to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.read()

        # Log audit event
        event = AuditEvent(
            matter_id=matter_id,
            event_type=AuditEventType.EXPORT_GENERATED,
            actor_id=actor_id,
            detail={"format": "docx", "size_bytes": len(docx_bytes)},
        )
        self.db.add(event)
        await self.db.commit()

        return docx_bytes

    async def _get_authoritative(self, model_class, matter_id: UUID):
        stmt = (
            select(model_class)
            .where(
                model_class.matter_id == matter_id,
                model_class.is_authoritative == True,
            )
            .order_by(desc(model_class.version_number))
            .limit(1)
        )
        result = await self.db.execute(stmt)
        return result.scalar_one_or_none()

    def _add_title_page(self, doc: DocxDocument, matter: Matter, brief: BriefVersion | None):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(matter.title or "Patent Application")
        run.bold = True
        run.font.size = Pt(24)

        doc.add_paragraph()  # spacer

        if matter.inventors:
            inventors_text = ", ".join(matter.inventors)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Inventors: {inventors_text}")
            run.font.size = Pt(12)

        if matter.assignee:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Assignee: {matter.assignee}")
            run.font.size = Pt(12)

        if brief and brief.structure_data:
            tech_field = brief.structure_data.get("technical_field", "")
            if tech_field:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"Technical Field: {tech_field}")
                run.font.size = Pt(11)

        doc.add_page_break()

    def _add_claims_section(self, doc: DocxDocument, claims: ClaimGraphVersion | None):
        doc.add_heading("Claims", level=1)

        if not claims or not claims.graph_data:
            doc.add_paragraph("No claims available.")
            doc.add_page_break()
            return

        nodes = claims.graph_data.get("nodes", [])
        for i, node in enumerate(nodes, 1):
            claim_type = node.get("type", "independent")
            text = node.get("text", "")
            deps = node.get("dependencies", [])

            p = doc.add_paragraph()
            run = p.add_run(f"{i}. ")
            run.bold = True

            if claim_type == "dependent" and deps:
                p.add_run(text)
            else:
                p.add_run(text)

            p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    def _add_spec_sections(self, doc: DocxDocument, spec: SpecVersion | None):
        if not spec or not spec.content_data:
            doc.add_heading("Specification", level=1)
            doc.add_paragraph("No specification available.")
            doc.add_page_break()
            return

        content = spec.content_data
        paragraphs = content.get("sections", [])

        # Group paragraphs by their section field
        from collections import defaultdict
        grouped: dict[str, list[dict]] = defaultdict(list)
        for para in paragraphs:
            grouped[para.get("section", "detailed_description")].append(para)

        # Map section types to patent headings, in order
        section_headings = [
            ("technical_field", "Technical Field"),
            ("background", "Background of the Invention"),
            ("summary", "Summary of the Invention"),
            ("brief_description_of_drawings", "Brief Description of the Drawings"),
            ("detailed_description", "Detailed Description of Preferred Embodiments"),
            ("definitions", "Definitions"),
            ("figure_descriptions", "Description of Figures"),
        ]

        for section_key, heading in section_headings:
            section_paras = grouped.get(section_key, [])
            if not section_paras:
                continue
            doc.add_heading(heading, level=1)
            for para in section_paras:
                text = para.get("text", "").strip()
                if text:
                    doc.add_paragraph(text)

        doc.add_page_break()

    def _add_abstract(self, doc: DocxDocument, spec: SpecVersion | None):
        doc.add_heading("Abstract", level=1)

        if spec and spec.content_data:
            paragraphs = spec.content_data.get("sections", [])
            abstract_texts = [
                p.get("text", "").strip()
                for p in paragraphs
                if p.get("section") == "abstract" and p.get("text", "").strip()
            ]
            if abstract_texts:
                doc.add_paragraph(" ".join(abstract_texts))
                return

        doc.add_paragraph("No abstract available.")
